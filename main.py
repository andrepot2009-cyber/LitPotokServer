import sys
import requests
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QMainWindow, QLabel, QVBoxLayout,
                             QHBoxLayout, QWidget, QPushButton, QStackedWidget,
                             QTextEdit, QListWidget, QMessageBox, QComboBox,
                             QLineEdit, QSlider, QListWidgetItem, QInputDialog)
from PyQt6.QtCore import Qt
from docx import Document

SERVER_URL = "https://litpotokserver-3.onrender.com"


class AuthWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ЛитПоток — Вход")
        self.resize(400, 320)
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout()
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.setSpacing(15)
        central.setLayout(layout)

        title = QLabel("Добро пожаловать в «ЛитПоток»!")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("font-size: 20px; font-weight: bold; color: #2c3e50;")
        layout.addWidget(title)

        self.username_input = QLineEdit()
        self.username_input.setPlaceholderText("Логин")
        self.username_input.setStyleSheet(
            "padding: 8px; font-size: 14px; border-radius: 4px; border: 1px solid #bdc3c7;")
        layout.addWidget(self.username_input)

        self.password_input = QLineEdit()
        self.password_input.setPlaceholderText("Пароль")
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.password_input.setStyleSheet(
            "padding: 8px; font-size: 14px; border-radius: 4px; border: 1px solid #bdc3c7;")
        layout.addWidget(self.password_input)

        role_layout = QHBoxLayout()
        role_layout.addWidget(QLabel("Роль:"))
        self.role_combo = QComboBox()
        self.role_combo.addItems(["Читатель", "Писатель"])
        self.role_combo.setStyleSheet("padding: 6px; font-size: 14px; border-radius: 4px;")
        role_layout.addWidget(self.role_combo)
        layout.addLayout(role_layout)

        btn_layout = QHBoxLayout()
        self.login_btn = QPushButton("Войти")
        self.login_btn.setStyleSheet(
            "font-size: 16px; padding: 10px 20px; background-color: #3498db; color: white; border-radius: 6px; border: none;")
        self.login_btn.clicked.connect(self.login)
        btn_layout.addWidget(self.login_btn)

        self.register_btn = QPushButton("Регистрация")
        self.register_btn.setStyleSheet(
            "font-size: 16px; padding: 10px 20px; background-color: #27ae60; color: white; border-radius: 6px; border: none;")
        self.register_btn.clicked.connect(self.register)
        btn_layout.addWidget(self.register_btn)
        layout.addLayout(btn_layout)

        self.status_label = QLabel("")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_label.setStyleSheet("color: #e74c3c; font-size: 13px;")
        layout.addWidget(self.status_label)

    def get_headers(self, token=None):
        headers = {"Content-Type": "application/json"}
        if token:
            headers["Authorization"] = f"Bearer {token}"
        return headers

    def login(self):
        username = self.username_input.text().strip()
        password = self.password_input.text().strip()
        if not username or not password:
            self.status_label.setText("Заполните поля")
            return

        # 🔐 Проверка длины пароля (bcrypt принимает ≤ 72 байт)
        if len(password.encode('utf-8')) > 72:
            self.status_label.setText("Пароль слишком длинный. Используйте до 72 байт (примерно 36 русских символов).")
            return

        try:
            resp = requests.post(f"{SERVER_URL}/auth/login", json={"username": username, "password": password},
                                 headers=self.get_headers(), timeout=10)
            if resp.status_code == 200:
                self.token = resp.json()["access_token"]
                self.role = "reader" if self.role_combo.currentText() == "Читатель" else "writer"
                self.close()
                self.main_app = MainWindow(username, self.role, self.token)
                self.main_app.show()
            else:
                self.status_label.setText(f"Ошибка: {resp.json().get('detail')}")
        except Exception as e:
            self.status_label.setText(f"Ошибка сети: {str(e)}")

    def register(self):
        username = self.username_input.text().strip()
        password = self.password_input.text().strip()
        if not username or not password:
            self.status_label.setText("Заполните поля")
            return

        # 🔐 Проверка длины пароля
        if len(password.encode('utf-8')) > 72:
            self.status_label.setText("Пароль слишком длинный. Используйте до 72 байт (примерно 36 русских символов).")
            return

        role = "reader" if self.role_combo.currentText() == "Читатель" else "writer"
        try:
            resp = requests.post(f"{SERVER_URL}/auth/register",
                                 json={"username": username, "password": password, "role": role},
                                 headers=self.get_headers(), timeout=10)
            if resp.status_code == 200:
                QMessageBox.information(self, "Успех", "Аккаунт создан! Войдите.")
                self.status_label.setText("")
            else:
                self.status_label.setText(f"Ошибка: {resp.json().get('detail')}")
        except Exception as e:
            self.status_label.setText(f"Ошибка сети: {str(e)}")


class MainWindow(QMainWindow):
    # ... весь остальной код MainWindow без изменений ...
    def __init__(self, username, role, token):
        super().__init__()
        self.username = username
        self.role = role
        self.token = token
        self.current_font_size = 15
        self.current_work_id = None
        self.current_work_author = None
        self.current_work_status = None
        self.feedback_data = []

        self.setWindowTitle(f"ЛитПоток — {username} ({'Читатель' if role == 'reader' else 'Писатель'})")
        self.resize(950, 650)

        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QHBoxLayout(main_widget)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        self.menu = QListWidget()
        self.menu.setFixedWidth(200)

        if self.role == "writer":
            self.menu.addItems(["📚 Библиотека", "📝 Черновики", "✏️ Редактор", "💬 Отзывы", "⚙️ Настройки"])
        else:
            self.menu.addItems(["📚 Библиотека", "📖 Чтение", "💬 Отзывы", "⚙️ Настройки"])

        self.menu.setStyleSheet("""
            QListWidget { background-color: #2c3e50; color: #ecf0f1; border: none; font-size: 15px; }
            QListWidget::item { padding: 12px; border-bottom: 1px solid #34495e; }
            QListWidget::item:selected { background-color: #3498db; color: white; }
            QListWidget::item:hover { background-color: #2980b9; }
        """)
        self.menu.currentRowChanged.connect(self.switch_page)
        main_layout.addWidget(self.menu)

        self.content_area = QStackedWidget()
        main_layout.addWidget(self.content_area)

        self.setup_library_page()

        if self.role == "writer":
            self.setup_drafts_page()

        self.setup_editor_page()
        self.setup_feedback_page()
        self.setup_settings_page()

        self.menu.setCurrentRow(0)

    def get_headers(self):
        return {"Content-Type": "application/json", "Authorization": f"Bearer {self.token}"}

    def switch_page(self, index):
        self.content_area.setCurrentIndex(index)
        if index == 0:
            self.refresh_library()
        elif self.role == "writer" and index == 1:
            self.refresh_drafts()

    def setup_library_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(15, 15, 15, 15)

        label = QLabel("📚 Все произведения" if self.role == "reader" else "📚 Опубликованные работы")
        label.setStyleSheet("font-size: 16px; font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        layout.addWidget(label)

        self.library_list = QListWidget()
        self.library_list.setStyleSheet("""
            QListWidget { font-size: 14px; background-color: #fafafa; color: #2c3e50; border: 1px solid #bdc3c7; border-radius: 4px; }
            QListWidget::item { padding: 8px; border-bottom: 1px solid #ecf0f1; }
            QListWidget::item:hover { background-color: #e8f6f3; }
            QListWidget::item:selected { background-color: #d5e8f7; color: #2c3e50; }
        """)
        self.library_list.itemDoubleClicked.connect(self.load_work)
        layout.addWidget(self.library_list)

        refresh_btn = QPushButton("🔄 Обновить")
        refresh_btn.setStyleSheet(
            "font-size: 14px; padding: 8px 16px; background-color: #3498db; color: white; border-radius: 4px; border: none;")
        refresh_btn.clicked.connect(self.refresh_library)
        layout.addWidget(refresh_btn)

        self.content_area.addWidget(page)

    def setup_drafts_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(15, 15, 15, 15)

        label = QLabel("📝 Мои черновики")
        label.setStyleSheet("font-size: 16px; font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        layout.addWidget(label)

        self.drafts_list = QListWidget()
        self.drafts_list.setStyleSheet("""
            QListWidget { font-size: 14px; background-color: #fff8e1; color: #2c3e50; border: 1px solid #fbc02d; border-radius: 4px; }
            QListWidget::item { padding: 8px; border-bottom: 1px solid #ffecb3; }
            QListWidget::item:hover { background-color: #ffecb3; }
            QListWidget::item:selected { background-color: #ffe082; color: #2c3e50; }
        """)
        self.drafts_list.itemDoubleClicked.connect(self.load_work)
        layout.addWidget(self.drafts_list)

        refresh_btn = QPushButton("🔄 Обновить")
        refresh_btn.setStyleSheet(
            "font-size: 14px; padding: 8px 16px; background-color: #f57c00; color: white; border-radius: 4px; border: none;")
        refresh_btn.clicked.connect(self.refresh_drafts)
        layout.addWidget(refresh_btn)

        self.content_area.addWidget(page)

    def setup_editor_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(15, 15, 15, 15)

        header_layout = QHBoxLayout()
        header_layout.setAlignment(Qt.AlignmentFlag.AlignRight)

        if self.role == "reader":
            self.review_btn = QPushButton("💬 Оставить отзыв")
            self.review_btn.setStyleSheet(
                "font-size: 13px; padding: 6px 12px; background-color: #9b59b6; color: white; border-radius: 4px; border: none;")
            self.review_btn.clicked.connect(self.open_feedback_modal)
            self.review_btn.setVisible(False)
            header_layout.addWidget(self.review_btn)

        layout.addLayout(header_layout)

        self.editor = QTextEdit()
        self.editor.setPlaceholderText("Текст произведения...")
        if self.role == "reader":
            self.editor.setReadOnly(True)
        self.editor.setStyleSheet(
            f"QTextEdit {{ font-size: {self.current_font_size}px; font-family: 'Segoe UI', 'Georgia', serif; padding: 10px; background-color: #fafafa; color: #2c3e50; border: 1px solid #bdc3c7; border-radius: 4px; }}")
        layout.addWidget(self.editor)

        btn_layout = QHBoxLayout()
        btn_layout.setAlignment(Qt.AlignmentFlag.AlignRight)

        if self.role == "writer":
            self.draft_btn = QPushButton("💾 Сохранить черновик")
            self.draft_btn.setStyleSheet(
                "font-size: 14px; padding: 8px 16px; background-color: #f39c12; color: white; border-radius: 4px; border: none;")
            self.draft_btn.clicked.connect(self.save_draft)
            self.draft_btn.setVisible(False)
            btn_layout.addWidget(self.draft_btn)

            self.publish_btn = QPushButton("🚀 Опубликовать")
            self.publish_btn.setStyleSheet(
                "font-size: 14px; padding: 8px 16px; background-color: #27ae60; color: white; border-radius: 4px; border: none;")
            self.publish_btn.clicked.connect(self.publish_work)
            self.publish_btn.setVisible(False)
            btn_layout.addWidget(self.publish_btn)

        layout.addLayout(btn_layout)
        self.content_area.addWidget(page)

    def setup_feedback_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(15, 15, 15, 15)

        self.feedback_header = QLabel("Выберите произведение в библиотеке")
        self.feedback_header.setStyleSheet("font-size: 15px; font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        layout.addWidget(self.feedback_header)

        form_layout = QHBoxLayout()
        self.category_combo = QComboBox()
        self.category_combo.addItems(["Сюжет", "Стиль", "Орфография", "Общее"])
        self.category_combo.setStyleSheet("padding: 6px; font-size: 14px;")

        self.feedback_input = QLineEdit()
        self.feedback_input.setPlaceholderText("Отзыв...")
        self.feedback_input.setStyleSheet(
            "padding: 6px; font-size: 14px; border: 1px solid #bdc3c7; border-radius: 4px;")

        self.add_feedback_btn = QPushButton("Отправить")
        self.add_feedback_btn.setStyleSheet(
            "font-size: 14px; padding: 6px 16px; background-color: #3498db; color: white; border-radius: 4px; border: none;")
        self.add_feedback_btn.clicked.connect(self.add_feedback)

        if self.role == "writer":
            self.category_combo.setVisible(False)
            self.feedback_input.setVisible(False)
            self.add_feedback_btn.setVisible(False)
            self.feedback_header.setText("💬 Отзывы читателей к вашим работам")

        form_layout.addWidget(self.category_combo)
        form_layout.addWidget(self.feedback_input, stretch=1)
        form_layout.addWidget(self.add_feedback_btn)
        layout.addLayout(form_layout)

        self.feedback_list = QListWidget()
        self.feedback_list.setWordWrap(True)
        self.feedback_list.setStyleSheet("""
            QListWidget { font-size: 14px; background-color: #fafafa; color: #2c3e50; border: 1px solid #bdc3c7; border-radius: 4px; }
            QListWidget::item { padding: 10px; border-bottom: 1px solid #ecf0f1; }
            QListWidget::item:hover { background-color: #e8f6f3; }
            QListWidget::item:selected { background-color: #d5e8f7; color: #2c3e50; }
        """)
        layout.addWidget(self.feedback_list)

        self.content_area.addWidget(page)

    def setup_settings_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        profile_label = QLabel(
            f"👤 Профиль: {self.username}\n🎭 Роль: {'Читатель' if self.role == 'reader' else 'Писатель'}")
        profile_label.setStyleSheet("font-size: 16px; font-weight: bold; color: #2c3e50; margin-bottom: 20px;")
        layout.addWidget(profile_label)

        font_label = QLabel("🔤 Размер шрифта:")
        font_label.setStyleSheet("font-size: 15px; font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        layout.addWidget(font_label)

        self.font_slider = QSlider(Qt.Orientation.Horizontal)
        self.font_slider.setRange(10, 30)
        self.font_slider.setValue(self.current_font_size)
        self.font_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.font_slider.setTickInterval(5)
        self.font_slider.valueChanged.connect(self.change_font_size)
        layout.addWidget(self.font_slider)

        layout.addSpacing(30)

        export_label = QLabel("📥 Экспорт:")
        export_label.setStyleSheet("font-size: 15px; font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        layout.addWidget(export_label)

        export_btn = QPushButton("📄 Сохранить как .docx")
        export_btn.setStyleSheet(
            "font-size: 15px; padding: 12px 24px; background-color: #8e44ad; color: white; border-radius: 4px; border: none;")
        export_btn.clicked.connect(self.export_to_docx)
        layout.addWidget(export_btn)

        layout.addStretch()
        self.content_area.addWidget(page)

    def refresh_library(self):
        try:
            resp = requests.get(f"{SERVER_URL}/works", params={"status": "published"}, headers=self.get_headers(),
                                timeout=10)
            if resp.status_code == 401:
                QMessageBox.warning(self, "Сессия", "Токен истёк. Войдите заново.")
                QApplication.quit()
                return
            resp.raise_for_status()
            works = resp.json()
            self.library_list.clear()
            if not works:
                msg = "Нет произведений" if self.role == "reader" else "Вы ничего не публиковали"
                item = QListWidgetItem(msg)
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEnabled)
                self.library_list.addItem(item)
            else:
                for w in works:
                    date_info = w.get('created_at', '...')
                    item = QListWidgetItem(f"{w['title']} — {w['author']} ({date_info})")
                    item.setData(Qt.ItemDataRole.UserRole, w['id'])
                    self.library_list.addItem(item)
        except Exception as e:
            self.library_list.clear()
            item = QListWidgetItem(f"Ошибка: {str(e)}")
            item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEnabled)
            self.library_list.addItem(item)

    def refresh_drafts(self):
        if self.role != "writer": return
        try:
            resp = requests.get(f"{SERVER_URL}/works", params={"status": "draft"}, headers=self.get_headers(),
                                timeout=10)
            resp.raise_for_status()
            works = resp.json()
            self.drafts_list.clear()
            if not works:
                item = QListWidgetItem("Нет черновиков")
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEnabled)
                self.drafts_list.addItem(item)
            else:
                for w in works:
                    date_info = w.get('created_at', '...')
                    item = QListWidgetItem(f"📝 {w['title']} ({date_info})")
                    item.setData(Qt.ItemDataRole.UserRole, w['id'])
                    self.drafts_list.addItem(item)
        except Exception as e:
            self.drafts_list.clear()
            item = QListWidgetItem(f"Ошибка: {str(e)}")
            item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEnabled)
            self.drafts_list.addItem(item)

    def load_work(self, item):
        work_id = item.data(Qt.ItemDataRole.UserRole)
        if not work_id: return

        try:
            resp = requests.get(f"{SERVER_URL}/works/{work_id}", headers=self.get_headers(), timeout=10)
            resp.raise_for_status()
            work = resp.json()

            self.editor.setPlainText(work['content'])
            self.current_work_id = work_id
            self.current_work_author = work.get('author', '')
            self.current_work_status = work.get('status', 'published')

            is_owner = (self.role == "writer" and self.current_work_author == self.username)

            if self.role == "reader":
                self.editor.setReadOnly(True)
            else:
                self.editor.setReadOnly(not is_owner)

            if self.role == "writer":
                is_draft = (self.current_work_status == "draft")
                self.draft_btn.setVisible(is_owner)
                self.publish_btn.setVisible(is_owner and is_draft)

            if self.role == "reader" and hasattr(self, 'review_btn'):
                self.review_btn.setVisible(True)

            self.load_feedback(work_id)

            if self.role == "writer":
                self.menu.setCurrentRow(2)
            else:
                self.menu.setCurrentRow(1)

            QMessageBox.information(self, "Успех", f"Открыто: {work['title']}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def open_feedback_modal(self):
        if not self.current_work_id or self.role != "reader":
            return

        text, ok = QInputDialog.getMultiLineText(self, "💬 Оставить отзыв", "Ваш отзыв:")
        if ok and text.strip():
            category, ok = QInputDialog.getItem(self, "Категория", "Выберите категорию:",
                                                ["Сюжет", "Стиль", "Орфография", "Общее"], 0, False)
            if ok:
                try:
                    resp = requests.post(f"{SERVER_URL}/feedback",
                                         json={"work_id": self.current_work_id, "category": category,
                                               "text": text.strip()},
                                         headers=self.get_headers(), timeout=10)
                    resp.raise_for_status()
                    QMessageBox.information(self, "Успех", "Отзыв отправлен!")
                    self.load_feedback(self.current_work_id)
                except Exception as e:
                    QMessageBox.critical(self, "Ошибка", str(e))

    def save_draft(self):
        if self.role != "writer": return

        title, ok = QInputDialog.getText(self, "Черновик", "Название черновика:")
        if not ok or not title.strip(): return

        text = self.editor.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "Внимание", "Текст пуст.")
            return

        try:
            if self.current_work_id and self.current_work_author == self.username:
                resp = requests.put(f"{SERVER_URL}/works/{self.current_work_id}",
                                    json={"title": title.strip(), "content": text, "status": "draft"},
                                    headers=self.get_headers(), timeout=10)
            else:
                resp = requests.post(f"{SERVER_URL}/works",
                                     json={"title": title.strip(), "content": text, "status": "draft"},
                                     headers=self.get_headers(), timeout=10)

            resp.raise_for_status()
            QMessageBox.information(self, "Успех", "Черновик сохранён!")
            self.refresh_drafts()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def publish_work(self):
        if self.role != "writer": return

        title, ok = QInputDialog.getText(self, "Публикация", "Название:")
        if not ok or not title.strip(): return

        text = self.editor.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "Внимание", "Текст пуст.")
            return

        try:
            if self.current_work_id and self.current_work_author == self.username:
                resp = requests.put(f"{SERVER_URL}/works/{self.current_work_id}",
                                    json={"title": title.strip(), "content": text, "status": "published"},
                                    headers=self.get_headers(), timeout=10)
            else:
                resp = requests.post(f"{SERVER_URL}/works",
                                     json={"title": title.strip(), "content": text, "status": "published"},
                                     headers=self.get_headers(), timeout=10)

            resp.raise_for_status()
            QMessageBox.information(self, "Успех", "Опубликовано!")
            self.editor.clear()
            self.current_work_id = None
            self.refresh_library()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def load_feedback(self, work_id):
        try:
            resp = requests.get(f"{SERVER_URL}/feedback/{work_id}", headers=self.get_headers(), timeout=10)
            resp.raise_for_status()
            self.feedback_data = resp.json()
            self.refresh_feedback_list()
        except Exception:
            self.feedback_data = []
            self.refresh_feedback_list()

    def add_feedback(self):
        if not self.current_work_id or self.role != "reader":
            QMessageBox.warning(self, "Внимание", "Выберите произведение.")
            return

        text = self.feedback_input.text().strip()
        if not text:
            QMessageBox.warning(self, "Внимание", "Введите текст.")
            return

        try:
            resp = requests.post(f"{SERVER_URL}/feedback",
                                 json={"work_id": self.current_work_id, "category": self.category_combo.currentText(),
                                       "text": text},
                                 headers=self.get_headers(), timeout=10)
            resp.raise_for_status()
            self.feedback_input.clear()
            self.load_feedback(self.current_work_id)
            QMessageBox.information(self, "Успех", "Отзыв отправлен!")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def refresh_feedback_list(self):
        self.feedback_list.clear()
        if not self.feedback_data:
            item = QListWidgetItem("Нет отзывов")
            item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEnabled)
            self.feedback_list.addItem(item)
            return

        for item_data in self.feedback_data:
            reviewer = item_data.get('reviewer', 'Аноним')
            date = item_data.get('created_at', '...')
            self.feedback_list.addItem(f"[{item_data['category']}] {item_data['text']}\n— {reviewer} | {date}")

    def change_font_size(self, size):
        self.current_font_size = size
        self.editor.setStyleSheet(
            f"QTextEdit {{ font-size: {size}px; font-family: 'Segoe UI', 'Georgia', serif; padding: 10px; background-color: #fafafa; color: #2c3e50; border: 1px solid #bdc3c7; border-radius: 4px; }}")

    def export_to_docx(self):
        text = self.editor.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "Внимание", "Текст пуст.")
            return

        try:
            doc = Document()
            doc.add_heading('ЛитПоток: Экспорт', 0)
            for p in text.split('\n'):
                doc.add_paragraph(p)
            doc.save(f"LitPotok_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            QMessageBox.information(self, "Успех", "Файл сохранён.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))


if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = AuthWindow()
    win.show()
    sys.exit(app.exec())